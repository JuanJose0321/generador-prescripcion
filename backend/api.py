"""
API para generación de documentos de prescripción fiscal.
Reemplaza valores en una plantilla Word preservando estilos y
verifica el resultado con Grok AI antes de responder.
"""
import base64
import io
import logging
import os
from datetime import datetime

from flask import Flask, jsonify, request, send_from_directory
from flask_cors import CORS
from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
import requests
from dotenv import load_dotenv

# Cargar variables de entorno
try:
    load_dotenv()
except:
    pass

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
)
logger = logging.getLogger("api")

# Variables para Grok
GROK_API_KEY = os.getenv("GROK_API_KEY")
GROK_API_URL = "https://api.x.ai/v1/messages"
GROK_MODEL = "grok-2-1212"
GROK_TIMEOUT = 20


# ==================== FUNCIONES DE VERIFICACIÓN ====================

def _extraer_texto(doc: Document) -> str:
    partes = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                partes.extend(p.text for p in cell.paragraphs)

    contenedores = [doc.element.body]
    for section in doc.sections:
        contenedores.append(section.header._element)
        contenedores.append(section.footer._element)
    for contenedor in contenedores:
        for caja in contenedor.findall(".//" + qn("w:txbxContent")):
            for p_element in caja.findall(".//" + qn("w:p")):
                partes.append(Paragraph(p_element, doc).text)

    return "\n".join(partes)


def _comparar_campos(texto_generado, reemplazos):
    ok, fallidos = [], []
    for valor_original, valor_nuevo in reemplazos.items():
        if valor_original == valor_nuevo:
            if valor_nuevo in texto_generado:
                ok.append(valor_original)
            else:
                fallidos.append({
                    "campo": valor_original,
                    "valor_esperado": valor_nuevo,
                    "razon": "El valor esperado no se encontró en el documento generado",
                })
            continue

        if valor_original in texto_generado:
            fallidos.append({
                "campo": valor_original,
                "valor_esperado": valor_nuevo,
                "razon": "El valor original todavía aparece en el documento generado",
            })
        elif valor_nuevo not in texto_generado:
            fallidos.append({
                "campo": valor_original,
                "valor_esperado": valor_nuevo,
                "razon": "El valor nuevo no se encontró en el documento generado (posible reemplazo parcial)",
            })
        else:
            ok.append(valor_original)
    return ok, fallidos


def _consultar_grok(texto_original, texto_generado, fallidos):
    if not GROK_API_KEY:
        logger.warning("GROK_API_KEY no configurada, se omite el análisis con IA")
        return "Análisis con IA omitido: no hay GROK_API_KEY configurada"

    prompt = (
        "Eres un asistente que revisa documentos legales generados a partir de una plantilla.\n"
        "Compara el TEXTO ORIGINAL (plantilla) contra el TEXTO GENERADO (documento final) y "
        "dime en 3-5 líneas si detectas valores de la plantilla que quedaron sin reemplazar, "
        "inconsistencias de formato, o cualquier otra anomalía. Sé breve y concreto.\n\n"
        f"--- TEXTO ORIGINAL ---\n{texto_original[:6000]}\n\n"
        f"--- TEXTO GENERADO ---\n{texto_generado[:6000]}\n\n"
        f"--- CAMPOS QUE FALLARON EN LA VERIFICACIÓN AUTOMÁTICA ---\n{fallidos}\n"
    )

    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {GROK_API_KEY}",
        "x-api-key": GROK_API_KEY,
        "anthropic-version": "2023-06-01",
    }
    payload = {
        "model": GROK_MODEL,
        "max_tokens": 500,
        "messages": [{"role": "user", "content": prompt}],
    }

    try:
        logger.info("Consultando Grok API (%s) para análisis semántico...", GROK_MODEL)
        resp = requests.post(GROK_API_URL, json=payload, headers=headers, timeout=GROK_TIMEOUT)
        resp.raise_for_status()
        data = resp.json()

        if isinstance(data.get("content"), list) and data["content"]:
            return data["content"][0].get("text", "").strip()
        if "choices" in data:
            return data["choices"][0]["message"]["content"].strip()

        logger.warning("Respuesta de Grok con formato inesperado: %s", data)
        return "Grok respondió en un formato inesperado; revisar logs del servidor."

    except requests.exceptions.Timeout:
        logger.error("Timeout consultando Grok API")
        return "Error: tiempo de espera agotado al consultar Grok API"
    except requests.exceptions.RequestException as exc:
        logger.error("Error consultando Grok API: %s", exc)
        return f"Error consultando Grok API: {exc}"
    except (KeyError, IndexError, ValueError) as exc:
        logger.error("Error interpretando respuesta de Grok: %s", exc)
        return f"Error interpretando la respuesta de Grok: {exc}"


def verificar_documento(template_path: str, documento_generado_bytes: bytes, reemplazos: dict) -> dict:
    try:
        doc_original = Document(template_path)
        doc_generado = Document(io.BytesIO(documento_generado_bytes))
    except Exception as exc:
        logger.exception("No se pudo abrir alguno de los documentos para verificar")
        return {
            "ok": [],
            "fallidos": [],
            "reportes": [f"No se pudo ejecutar la verificación: {exc}"],
        }

    texto_original = _extraer_texto(doc_original)
    texto_generado = _extraer_texto(doc_generado)

    ok, fallidos = _comparar_campos(texto_generado, reemplazos)

    reportes = []
    if fallidos:
        reportes.append(f"{len(fallidos)} valor(es) no se reemplazaron correctamente.")
    else:
        reportes.append("Todos los valores se reemplazaron correctamente según la verificación automática.")

    analisis_ia = _consultar_grok(texto_original, texto_generado, fallidos)
    reportes.append(analisis_ia)

    return {"ok": ok, "fallidos": fallidos, "reportes": reportes}


# ==================== FIN FUNCIONES DE VERIFICACIÓN ====================

app = Flask(__name__)
CORS(app)

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
TEMPLATE_PATH = os.path.join(BASE_DIR, "FORMATO_ORIGINAL_2.docx")

CAMPOS_REQUERIDOS = [
    "representante", "claveAnterior", "claveActual", "propietario",
    "predialInicio", "serviciosInicio", "fechaPeticion",
    "folioExpediente1", "folioExpediente2", "fechaEjecucion",
    "caracterSolicitante",
]

_CARACTER_SOLICITANTE = {
    "PROPIETARIO": "EN SU CARÁCTER DE PROPIETARIO.",
    "IMPUESTOS": "EN SU CARÁCTER DE INTERESADO EN CUBRIR LOS ADEUDOS FISCALES POR CONCEPTO DE IMPUESTO PREDIAL Y SERVICIOS PÚBLICOS MUNICIPALES.",
}

_MESES = {
    1: "enero", 2: "febrero", 3: "marzo", 4: "abril", 5: "mayo", 6: "junio",
    7: "julio", 8: "agosto", 9: "septiembre", 10: "octubre", 11: "noviembre", 12: "diciembre",
}

_UNIDADES = ["", "uno", "dos", "tres", "cuatro", "cinco", "seis", "siete", "ocho", "nueve"]
_DIEZ_A_VEINTINUEVE = {
    10: "diez", 11: "once", 12: "doce", 13: "trece", 14: "catorce", 15: "quince",
    16: "dieciséis", 17: "diecisiete", 18: "dieciocho", 19: "diecinueve",
    20: "veinte", 21: "veintiuno", 22: "veintidós", 23: "veintitrés", 24: "veinticuatro",
    25: "veinticinco", 26: "veintiséis", 27: "veintisiete", 28: "veintiocho", 29: "veintinueve",
}
_DECENAS = {
    30: "treinta", 40: "cuarenta", 50: "cincuenta", 60: "sesenta",
    70: "setenta", 80: "ochenta", 90: "noventa",
}
_CENTENAS = {
    100: "cien", 200: "doscientos", 300: "trescientos", 400: "cuatrocientos", 500: "quinientos",
    600: "seiscientos", 700: "setecientos", 800: "ochocientos", 900: "novecientos",
}


def _menores_a_cien_en_palabras(n):
    if n == 0:
        return ""
    if n < 10:
        return _UNIDADES[n]
    if n <= 29:
        return _DIEZ_A_VEINTINUEVE[n]
    decena, unidad = (n // 10) * 10, n % 10
    if unidad == 0:
        return _DECENAS[decena]
    return f"{_DECENAS[decena]} y {_UNIDADES[unidad]}"


def _menores_a_mil_en_palabras(n):
    if n == 0:
        return ""
    centena, resto = (n // 100) * 100, n % 100
    if centena == 0:
        return _menores_a_cien_en_palabras(resto)
    texto_centena = "ciento" if (centena == 100 and resto > 0) else _CENTENAS[centena]
    if resto == 0:
        return texto_centena
    return f"{texto_centena} {_menores_a_cien_en_palabras(resto)}"


def _numero_a_palabras(n):
    if n == 0:
        return "cero"
    miles, resto = n // 1000, n % 1000
    partes = []
    if miles > 0:
        partes.append("mil" if miles == 1 else f"{_menores_a_mil_en_palabras(miles)} mil")
    if resto > 0:
        partes.append(_menores_a_mil_en_palabras(resto))
    return " ".join(partes)


def convertir_numero_a_texto_año(periodo):
    """
    Convierte el año de un período "bimestre/año" a su forma escrita en
    español. Ej: "1/2008" -> "dos mil ocho", "1/1998" -> "mil novecientos
    noventa y ocho", "6/2024" -> "dos mil veinticuatro".
    """
    año = int(periodo.split("/")[-1].strip())
    return _numero_a_palabras(año)


def validar_folio(folio, nombre_campo):
    """
    Valida que un folio de expediente contenga solo dígitos. Se usa
    para folioExpediente1/2, que se insertan tal cual (sin conversión)
    en el número de expediente.
    """
    if not folio.isdigit():
        raise ValueError(f"{nombre_campo} debe contener solo números, recibido: {folio!r}")


def convertir_fecha_a_palabras(fecha_iso):
    """
    Convierte una fecha ISO "AAAA-MM-DD" a la forma escrita que usa la
    plantilla, ej. "2026-07-13" -> "trece días del mes de julio del año
    dos mil veintiséis".
    """
    fecha = datetime.strptime(fecha_iso, "%Y-%m-%d")
    dia_texto = _numero_a_palabras(fecha.day)
    mes_texto = _MESES[fecha.month]
    año_texto = _numero_a_palabras(fecha.year)
    return f"{dia_texto} días del mes de {mes_texto} del año {año_texto}"


def convertir_fecha_a_formato_dia(fecha_iso):
    """
    Convierte una fecha ISO "AAAA-MM-DD" al formato "el día DD de mes de
    AAAA" usado en el párrafo 11, ej. "2026-09-28" -> "el día 28 de
    septiembre de 2026".
    """
    fecha = datetime.strptime(fecha_iso, "%Y-%m-%d")
    mes_texto = _MESES[fecha.month]
    return f"el día {fecha.day:02d} de {mes_texto} de {fecha.year}"


def _recolectar_parrafos(doc):
    """
    Reúne TODOS los párrafos del documento sobre los que puede haber que
    reemplazar texto o limpiar resaltado: párrafos del cuerpo, celdas de
    tabla, y párrafos dentro de cuadros de texto (encabezado, pie de
    página y cuerpo). Los cuadros de texto viven en XML como
    `w:txbxContent` y `python-docx` no los expone vía `doc.paragraphs`,
    así que hay que ubicarlos a mano y envolverlos con `Paragraph`.
    """
    parrafos = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parrafos.extend(cell.paragraphs)

    vistos = {id(p._p) for p in parrafos}

    contenedores = [doc.element.body]
    for section in doc.sections:
        contenedores.append(section.header._element)
        contenedores.append(section.footer._element)

    for contenedor in contenedores:
        for caja in contenedor.findall(".//" + qn("w:txbxContent")):
            for p_element in caja.findall(".//" + qn("w:p")):
                if id(p_element) not in vistos:
                    vistos.add(id(p_element))
                    parrafos.append(Paragraph(p_element, doc))

    return parrafos


def limpiar_resaltado(doc):
    """
    Quita cualquier resaltado (ej. amarillo) de todos los runs del
    documento, incluidos los que viven dentro de cuadros de texto,
    encabezado y pie de página.
    """
    total = 0
    for paragraph in _recolectar_parrafos(doc):
        for run in paragraph.runs:
            if run.font.highlight_color is not None:
                run.font.highlight_color = None
                total += 1
    logger.info("Resaltado eliminado de %d run(s)", total)


def _aplicar_reemplazo(paragraph, pos_start, pos_end, reemplazar):
    """
    Escribe el texto de reemplazo sobre los runs de `paragraph` tocados
    por el rango [pos_start, pos_end) del texto concatenado del párrafo.

    Nunca modifica runs de longitud cero (`run.text == ""`): son los que
    contienen marcas internas de campos de Word (`w:fldChar`,
    `w:instrText`, típicos de campos REF usados para reutilizar un
    valor ya capturado en otra parte del documento). Escribir en ellos
    —incluso vaciarlos— dispara `clear_content()` en python-docx, que
    borra esos elementos y corrompe la estructura del campo (Word puede
    terminar mostrando el código crudo, ej. "REF periodoadeudosm  \\*
    MERGEFORMAT", en vez del valor). Al excluirlos, su XML queda
    intacto y solo se tocan los runs con texto real.

    `reemplazar` puede ser un string (comportamiento simple: todo el
    texto va al primer run tocado y el resto se vacía) o una lista de
    fragmentos. Con una lista, los runs tocados se agrupan en tramos
    contiguos que comparten el mismo estado de negrita, y cada
    fragmento se asigna a un tramo en orden — así un valor como
    "(1/2011)" que en la plantilla está en negrita dentro de un texto
    que no lo está conserva esa negrita en el resultado, en vez de que
    todo el reemplazo termine con el formato de un único run "ganador".
    Si la cantidad de fragmentos no coincide con la de tramos, se cae
    al comportamiento simple para no perder texto.
    """
    tocados = []
    pos_run = 0
    for run in paragraph.runs:
        run_start = pos_run
        run_end = pos_run + len(run.text)
        if len(run.text) > 0 and run_start < pos_end and run_end > pos_start:
            tocados.append((run, run_start, run_end))
        pos_run = run_end

    if not tocados:
        return

    segmentos = list(reemplazar) if isinstance(reemplazar, (list, tuple)) else [reemplazar]

    tramos = []
    for run, run_start, run_end in tocados:
        bold = run.font.bold
        if tramos and tramos[-1][0] == bold:
            tramos[-1][1].append((run, run_start, run_end))
        else:
            tramos.append((bold, [(run, run_start, run_end)]))

    if len(segmentos) != len(tramos):
        segmentos = ["".join(segmentos)]
        tramos = [(None, tocados)]

    for (_bold, runs_tramo), texto_nuevo in zip(tramos, segmentos):
        primero = True
        for run, run_start, run_end in runs_tramo:
            offset_start = max(0, pos_start - run_start)
            offset_end = min(len(run.text), pos_end - run_start)
            texto_run = texto_nuevo if primero else ""
            primero = False
            run.text = run.text[:offset_start] + texto_run + run.text[offset_end:]


def reemplazar_inteligente(doc, reemplazos):
    """
    Reemplaza valores en el documento SIN reconstruir párrafos.
    Preserva estilos (bold, tamaño, fuente) porque nunca borra ni
    recrea `run`s, solo modifica el texto de cada uno en su lugar (ver
    `_aplicar_reemplazo` para los detalles de negrita y campos de Word).

    Maneja valores fragmentados entre varios runs, por ejemplo un run
    con "1/20" seguido de otro run con "03" que juntos forman "1/2003".

    Cada párrafo se procesa de forma independiente: un mismo valor que
    se repite en varios párrafos (frecuente en documentos legales que
    citan la misma clave o fecha en distintas cláusulas) se reemplaza en
    todas sus apariciones, no solo en la primera del documento.

    Las coincidencias se ubican sobre el texto ORIGINAL e inmutable del
    párrafo (no sobre el texto ya modificado), para que el valor recién
    insertado de un campo no sea "recapturado" por la búsqueda de otro
    campo cuyo valor nuevo coincida por casualidad con esa cadena.
    """
    contador = {k: 0 for k in reemplazos}

    def procesar_parrafos(parrafos):
        for paragraph in parrafos:
            if not paragraph.runs:
                continue

            texto_original = "".join(r.text for r in paragraph.runs)
            if not texto_original:
                continue

            matches = []
            pos = 0
            while pos < len(texto_original):
                mejor = None
                for buscar, reemplazar in reemplazos.items():
                    idx = texto_original.find(buscar, pos)
                    if idx != -1 and (mejor is None or idx < mejor[0]):
                        mejor = (idx, buscar, reemplazar)
                if mejor is None:
                    break
                idx, buscar, reemplazar = mejor
                matches.append((idx, idx + len(buscar), buscar, reemplazar))
                pos = idx + len(buscar)

            if not matches:
                continue

            # Se aplica de derecha a izquierda para que los offsets ya
            # calculados para los matches restantes sigan siendo válidos.
            for pos_start, pos_end, buscar, reemplazar in reversed(matches):
                _aplicar_reemplazo(paragraph, pos_start, pos_end, reemplazar)
                contador[buscar] += 1

    procesar_parrafos(_recolectar_parrafos(doc))

    ok = sum(1 for c in contador.values() if c > 0)
    logger.info("Reemplazo completado: %d/%d campos encontrados", ok, len(contador))
    for buscar, veces in contador.items():
        if veces == 0:
            logger.warning("No se encontró el valor a reemplazar: %r", buscar)

    return contador


def _clausula_periodo(conector, año_fin_texto, periodo_fin, concepto,
                       año_inicio_texto, periodo_inicio):
    """
    Construye el par (buscar, reemplazar) de una cláusula de período
    ("primer bimestre ... al/hasta el sexto bimestre ..., por concepto
    de ..."). Solo el INICIO del período es dinámico; el FIN
    (año_fin_texto/periodo_fin) es el mismo texto fijo de la plantilla
    tanto en la búsqueda como en el reemplazo, así que nunca cambia.

    La misma frase de inicio ("dos mil once (1/2011)") se repite dos
    veces sin cambios dentro de un mismo párrafo (una vez para predial y
    otra para servicios) e incluso con texto idéntico entre
    CUARTO/QUINTO/SEGUNDO, así que el sufijo "por concepto de
    <concepto>" se incluye en la clave para que cada ocurrencia sea
    textualmente única en todo el documento y no se resuelvan todas al
    mismo valor.

    En la plantilla, cada paréntesis "(año/período)" está en negrita y
    el resto de la frase no. `reemplazar` se arma como un ciclo de
    fragmentos [texto, "(periodo)", texto, "(periodo)", texto] que
    `_aplicar_reemplazo` empareja con los tramos de negrita del
    original, así el paréntesis nuevo también queda en negrita en vez
    de heredar el formato plano del texto que lo rodea.
    """
    buscar = (
        f"primer bimestre del ejercicio fiscal dos mil once (1/2011), "
        f"{conector} sexto bimestre del ejercicio fiscal {año_fin_texto} ({periodo_fin}), "
        f"por concepto de {concepto}"
    )
    reemplazar = [
        f"primer bimestre del ejercicio fiscal {año_inicio_texto} ",
        f"({periodo_inicio}), ",
        f"{conector} sexto bimestre del ejercicio fiscal {año_fin_texto} ",
        f"({periodo_fin}), ",
        f"por concepto de {concepto}",
    ]
    return buscar, reemplazar


def construir_reemplazos(datos):
    predial_inicio_texto = convertir_numero_a_texto_año(datos["predialInicio"])
    servicios_inicio_texto = convertir_numero_a_texto_año(datos["serviciosInicio"])

    fecha_ejecucion_texto = convertir_fecha_a_palabras(datos["fechaEjecucion"])
    fecha_peticion_texto = convertir_fecha_a_formato_dia(datos["fechaPeticion"])

    validar_folio(datos["folioExpediente1"], "folioExpediente1")
    validar_folio(datos["folioExpediente2"], "folioExpediente2")

    caracter = datos["caracterSolicitante"]
    if caracter not in _CARACTER_SOLICITANTE:
        opciones = ", ".join(_CARACTER_SOLICITANTE)
        raise ValueError(f"caracterSolicitante debe ser uno de: {opciones} (recibido: {caracter!r})")

    reemplazos = {
        # Párrafo 8 (destinatario): mismo nombre de plantilla "PEDRO
        # PÉREZ FLORES" que en el cuadro SOLICITANTE, ambos con el
        # valor de representante.
        "C. PEDRO PÉREZ FLORES.": f"C. {datos['representante']}.",

        # Párrafo 25 (suscriptor de la petición) y párrafo 37 (frase de
        # "suscrito por..." agregada al RESUELVE): mismo contexto,
        # también representante.
        "el C. PEDRO PÉREZ FLORES, en su carácter de interesado":
            f"el C. {datos['representante']}, en su carácter de interesado",

        # Párrafo 15 y 16 (claves catastrales)
        "1100 09 119 041": datos["claveAnterior"],
        "1109 09 119 041": datos["claveActual"],

        # Párrafo 37 (RESUELVE SEGUNDO): mismo nombre en su calidad de
        # propietario. Ahí está en negrita (campo CONTRIBUYENTE)
        # mientras que "a nombre de " no lo está, así que se manda como
        # dos fragmentos para que la negrita del nombre se conserve.
        "a nombre de PEDRO PÉREZ FLORES,": ["a nombre de ", f"{datos['propietario']},"],

        # Cuadro de texto SOLICITANTE: nombre y "carácter" son ahora
        # dos campos independientes (el carácter viene de un dropdown),
        # así que se buscan por separado en vez de como una sola frase.
        "PEDRO PÉREZ FLORES, ": f"{datos['representante']}, ",
        "EN SU CARÁCTER DE PROPIETARIO.": _CARACTER_SOLICITANTE[caracter],

        # Cuadro de texto EXPEDIENTE (aparece dos veces en el documento).
        # El año queda fijo en 2026; solo cambian los dos folios.
        "TM/PRESCRIPCIÓN/07/01/2026":
            f"TM/PRESCRIPCIÓN/{datos['folioExpediente1']}/{datos['folioExpediente2']}/2026",

        # Párrafo 45 (fecha de ejecución), frase completa incluido el año
        "trece días del mes de julio del año dos mil veintiséis": fecha_ejecucion_texto,

        # Párrafo 11 (fecha de petición recibida)
        "el día 03 de julio de 2026": fecha_peticion_texto,
    }

    # Párrafo 29 (CUARTO), conector "hasta el", fin fijo 6/2025
    for concepto, año_i, per_i in (
        ("Impuesto Predial", predial_inicio_texto, datos["predialInicio"]),
        ("Servicios Públicos Municipales", servicios_inicio_texto, datos["serviciosInicio"]),
    ):
        buscar, reemplazar = _clausula_periodo("hasta el", "dos mil veinticinco", "6/2025", concepto, año_i, per_i)
        reemplazos[buscar] = reemplazar

    # Párrafo 31 (QUINTO), conector "al", fin fijo 6/2025
    for concepto, año_i, per_i in (
        ("Impuesto Predial", predial_inicio_texto, datos["predialInicio"]),
        ("Servicios Públicos Municipales", servicios_inicio_texto, datos["serviciosInicio"]),
    ):
        buscar, reemplazar = _clausula_periodo("al", "dos mil veinticinco", "6/2025", concepto, año_i, per_i)
        reemplazos[buscar] = reemplazar

    # Párrafo 37 (RESUELVE SEGUNDO), conector "al", fin fijo 6/2020
    for concepto, año_i, per_i in (
        ("Impuesto Predial", predial_inicio_texto, datos["predialInicio"]),
        ("Servicios Públicos Municipales", servicios_inicio_texto, datos["serviciosInicio"]),
    ):
        buscar, reemplazar = _clausula_periodo("al", "dos mil veinte", "6/2020", concepto, año_i, per_i)
        reemplazos[buscar] = reemplazar

    return reemplazos


@app.route("/api/generar", methods=["POST"])
def generar_documento():
    try:
        datos = request.json or {}

        faltantes = [c for c in CAMPOS_REQUERIDOS if not datos.get(c)]
        if faltantes:
            logger.warning("Solicitud incompleta, faltan: %s", faltantes)
            return jsonify({"error": f"Faltan campos: {', '.join(faltantes)}"}), 400

        if not os.path.exists(TEMPLATE_PATH):
            logger.error("No se encontró la plantilla en %s", TEMPLATE_PATH)
            return jsonify({"error": "No se encontró FORMATO_ORIGINAL_2.docx"}), 500

        logger.info("Generando documento para representante=%s", datos["representante"])

        try:
            reemplazos = construir_reemplazos(datos)
        except (ValueError, IndexError) as exc:
            logger.warning("Formato de campo inválido: %s", exc)
            return jsonify({"error": f"Formato de campo inválido: {exc}"}), 400

        doc = Document(TEMPLATE_PATH)
        limpiar_resaltado(doc)
        reemplazar_inteligente(doc, reemplazos)

        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_bytes = doc_io.getvalue()

        logger.info("Ejecutando verificación con IA...")
        # verificar_documento compara valores de texto plano; algunos
        # reemplazos vienen como lista de fragmentos (para preservar
        # negrita en `reemplazar_inteligente`), así que se aplanan aquí.
        reemplazos_verificacion = {
            buscar: ("".join(valor) if isinstance(valor, (list, tuple)) else valor)
            for buscar, valor in reemplazos.items()
        }
        verificacion = verificar_documento(TEMPLATE_PATH, doc_bytes, reemplazos_verificacion)
        logger.info(
            "Verificación completa: %d ok, %d fallidos",
            len(verificacion.get("ok", [])),
            len(verificacion.get("fallidos", [])),
        )

        filename = f'Prescripcion_{datos["representante"].split()[0]}_{datetime.now().strftime("%Y%m%d")}.docx'

        return jsonify({
            "filename": filename,
            "mimetype": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "documento_base64": base64.b64encode(doc_bytes).decode("ascii"),
            "verificacion": verificacion,
        })

    except Exception as exc:
        logger.exception("Error generando documento")
        return jsonify({"error": str(exc)}), 500


@app.route("/", methods=["GET"])
def index():
    try:
        # En Vercel, el archivo está en backend/../frontend
        # En local, también está en backend/../frontend
        script_dir = os.path.dirname(os.path.abspath(__file__))
        project_root = os.path.dirname(script_dir)
        frontend_dir = os.path.join(project_root, "frontend")

        logger.info("Buscando index.html en: %s", frontend_dir)

        if not os.path.exists(frontend_dir):
            logger.error("Directorio frontend no existe: %s", frontend_dir)
            return jsonify({"error": "Frontend directory not found"}), 500

        index_path = os.path.join(frontend_dir, "index.html")
        if not os.path.exists(index_path):
            logger.error("index.html no existe: %s", index_path)
            return jsonify({"error": "index.html not found"}), 500

        return send_from_directory(frontend_dir, "index.html")
    except Exception as e:
        logger.error("Error sirviendo index.html: %s", e)
        return jsonify({"error": f"Error: {str(e)}"}), 500


if __name__ == "__main__":
    app.run(debug=True)
