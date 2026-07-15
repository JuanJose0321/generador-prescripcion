"""
Verificación posterior a la generación del documento.
Compara el documento original (plantilla) contra el generado para
detectar valores que no se reemplazaron, y usa Grok (x.ai) para un
análisis semántico adicional.
"""
import io
import logging
import os

import requests
from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from dotenv import load_dotenv

load_dotenv()

logger = logging.getLogger("verificador_ia")

GROK_API_KEY = os.getenv("GROK_API_KEY")
GROK_API_URL = "https://api.x.ai/v1/messages"
GROK_MODEL = "grok-2-1212"
GROK_TIMEOUT = 20


def _extraer_texto(doc: Document) -> str:
    partes = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                partes.extend(p.text for p in cell.paragraphs)

    # Los cuadros de texto (ej. el bloque EXPEDIENTE/SOLICITANTE) viven
    # en XML como w:txbxContent y no aparecen en doc.paragraphs.
    contenedores = [doc.element.body]
    for section in doc.sections:
        contenedores.append(section.header._element)
        contenedores.append(section.footer._element)
    for contenedor in contenedores:
        for caja in contenedor.findall(".//" + qn("w:txbxContent")):
            for p_element in caja.findall(".//" + qn("w:p")):
                partes.append(Paragraph(p_element, doc).text)

    return "\n".join(partes)


def _comparar_campos(texto_generado, reemplazos):
    """Detecta, campo por campo, si el valor original sigue presente
    (no se reemplazó) o si el valor nuevo no aparece (reemplazo parcial)."""
    ok, fallidos = [], []
    for valor_original, valor_nuevo in reemplazos.items():
        # Algunos campos (ej. una opción de dropdown que coincide con el
        # valor por defecto de la plantilla) no cambian el texto. En ese
        # caso "el valor original sigue presente" es el resultado
        # correcto, no una falla.
        if valor_original == valor_nuevo:
            if valor_nuevo in texto_generado:
                ok.append(valor_original)
            else:
                fallidos.append({
                    "campo": valor_original,
                    "valor_esperado": valor_nuevo,
                    "razon": "El valor esperado no se encontró en el documento generado",
                })
            continue

        if valor_original in texto_generado:
            fallidos.append({
                "campo": valor_original,
                "valor_esperado": valor_nuevo,
                "razon": "El valor original todavía aparece en el documento generado",
            })
        elif valor_nuevo not in texto_generado:
            fallidos.append({
                "campo": valor_original,
                "valor_esperado": valor_nuevo,
                "razon": "El valor nuevo no se encontró en el documento generado (posible reemplazo parcial)",
            })
        else:
            ok.append(valor_original)
    return ok, fallidos


def _consultar_grok(texto_original, texto_generado, fallidos):
    if not GROK_API_KEY:
        logger.warning("GROK_API_KEY no configurada, se omite el análisis con IA")
        return "Análisis con IA omitido: no hay GROK_API_KEY configurada en .env"

    prompt = (
        "Eres un asistente que revisa documentos legales generados a partir de una plantilla.\n"
        "Compara el TEXTO ORIGINAL (plantilla) contra el TEXTO GENERADO (documento final) y "
        "dime en 3-5 líneas si detectas valores de la plantilla que quedaron sin reemplazar, "
        "inconsistencias de formato, o cualquier otra anomalía. Sé breve y concreto.\n\n"
        f"--- TEXTO ORIGINAL ---\n{texto_original[:6000]}\n\n"
        f"--- TEXTO GENERADO ---\n{texto_generado[:6000]}\n\n"
        f"--- CAMPOS QUE FALLARON EN LA VERIFICACIÓN AUTOMÁTICA ---\n{fallidos}\n"
    )

    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {GROK_API_KEY}",
        "x-api-key": GROK_API_KEY,
        "anthropic-version": "2023-06-01",
    }
    payload = {
        "model": GROK_MODEL,
        "max_tokens": 500,
        "messages": [{"role": "user", "content": prompt}],
    }

    try:
        logger.info("Consultando Grok API (%s) para análisis semántico...", GROK_MODEL)
        resp = requests.post(GROK_API_URL, json=payload, headers=headers, timeout=GROK_TIMEOUT)
        resp.raise_for_status()
        data = resp.json()

        if isinstance(data.get("content"), list) and data["content"]:
            return data["content"][0].get("text", "").strip()
        if "choices" in data:
            return data["choices"][0]["message"]["content"].strip()

        logger.warning("Respuesta de Grok con formato inesperado: %s", data)
        return "Grok respondió en un formato inesperado; revisar logs del servidor."

    except requests.exceptions.Timeout:
        logger.error("Timeout consultando Grok API")
        return "Error: tiempo de espera agotado al consultar Grok API"
    except requests.exceptions.RequestException as exc:
        logger.error("Error consultando Grok API: %s", exc)
        return f"Error consultando Grok API: {exc}"
    except (KeyError, IndexError, ValueError) as exc:
        logger.error("Error interpretando respuesta de Grok: %s", exc)
        return f"Error interpretando la respuesta de Grok: {exc}"


def verificar_documento(template_path: str, documento_generado_bytes: bytes, reemplazos: dict) -> dict:
    """
    Compara la plantilla original contra el documento generado.

    Retorna: {"ok": [...], "fallidos": [...], "reportes": [...]}
    """
    try:
        doc_original = Document(template_path)
        doc_generado = Document(io.BytesIO(documento_generado_bytes))
    except Exception as exc:
        logger.exception("No se pudo abrir alguno de los documentos para verificar")
        return {
            "ok": [],
            "fallidos": [],
            "reportes": [f"No se pudo ejecutar la verificación: {exc}"],
        }

    texto_original = _extraer_texto(doc_original)
    texto_generado = _extraer_texto(doc_generado)

    ok, fallidos = _comparar_campos(texto_generado, reemplazos)

    reportes = []
    if fallidos:
        reportes.append(f"{len(fallidos)} valor(es) no se reemplazaron correctamente.")
    else:
        reportes.append("Todos los valores se reemplazaron correctamente según la verificación automática.")

    analisis_ia = _consultar_grok(texto_original, texto_generado, fallidos)
    reportes.append(analisis_ia)

    return {"ok": ok, "fallidos": fallidos, "reportes": reportes}
