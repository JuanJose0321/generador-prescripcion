from flask import Flask, request, send_file
from flask_cors import CORS
from docx import Document
import io
from datetime import datetime
import os

app = Flask(__name__)
CORS(app)

def reemplazar_inteligente_final(doc, reemplazos):
    """
    Reemplaza valores SIN NUNCA reconstruir los párrafos.
    Maneja fragmentos identificando los runs exactos que los contienen.
    """
    
    contador = {k: 0 for k in reemplazos.keys()}
    
    def procesar_parrafos(parrafos):
        for para_idx, paragraph in enumerate(parrafos):
            # PASO 1: Intentar reemplazo directo en cada run
            for run in paragraph.runs:
                for buscar, reemplazar in reemplazos.items():
                    if buscar in run.text:
                        ocurrencias = run.text.count(buscar)
                        run.text = run.text.replace(buscar, reemplazar)
                        contador[buscar] += ocurrencias
            
            # PASO 2: Buscar valores fragmentados (sin reconstruir)
            for buscar, reemplazar in reemplazos.items():
                if contador[buscar] > 0:
                    # Ya se reemplazó
                    continue
                
                # Juntar texto para buscar
                texto_completo = ''.join([r.text for r in paragraph.runs])
                
                if buscar not in texto_completo:
                    # No está ni completo ni fragmentado
                    continue
                
                # ESTÁ FRAGMENTADO - Reemplazarlo sin reconstruir
                # Estrategia: mapear posición del texto en runs
                
                pos_en_completo = 0
                pos_start = texto_completo.find(buscar)
                pos_end = pos_start + len(buscar)
                
                # Identificar cuáles runs contienen este rango
                pos_run = 0
                for run in paragraph.runs:
                    run_start = pos_run
                    run_end = pos_run + len(run.text)
                    
                    # ¿Este run intersecta con el valor a reemplazar?
                    if run_start < pos_end and run_end > pos_start:
                        # Sí - reemplazar esta parte
                        offset_start = max(0, pos_start - run_start)
                        offset_end = min(len(run.text), pos_end - run_start)
                        
                        # Parte del valor original que está en este run
                        valor_en_run = run.text[offset_start:offset_end]
                        
                        # Índice en el valor original
                        idx_en_valor = run_start + offset_start - pos_start
                        
                        # Parte del reemplazo correspondiente
                        reemplazo_para_run = reemplazar[idx_en_valor:idx_en_valor + len(valor_en_run)]
                        
                        # Reemplazar en el run
                        run.text = run.text[:offset_start] + reemplazo_para_run + run.text[offset_end:]
                    
                    pos_run = run_end
                
                contador[buscar] += 1
    
    print("\n" + "="*80)
    print("REEMPLAZANDO (SIN reconstruir párrafos)")
    print("="*80 + "\n")
    
    procesar_parrafos(doc.paragraphs)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                procesar_parrafos(cell.paragraphs)
    
    print("\n" + "="*80)
    print("RESULTADO:")
    print("="*80)
    ok = sum(1 for c in contador.values() if c > 0)
    print(f"✅ {ok}/{len(contador)} campos reemplazados\n")

@app.route('/api/generar', methods=['POST'])
def generar_documento():
    try:
        datos = request.json
        
        campos = [
            'fechaExpedicion', 'representante', 'institucion', 'fechaPeticion',
            'claveAnterior', 'claveActual', 'propietario',
            'predialInicio', 'predialFin', 'serviciosInicio', 'serviciosFin',
            'fechaEjecucion'
        ]
        
        for campo in campos:
            if campo not in datos or not datos[campo]:
                return {'error': f'Falta: {campo}'}, 400
        
        template_path = 'FORMATO_ORIGINAL.docx'
        if not os.path.exists(template_path):
            return {'error': 'No se encontró FORMATO_ORIGINAL.docx'}, 500
        
        doc = Document(template_path)
        
        reemplazos = {
            'RAFAEL BERMÚDEZ GUTIÉRREZ': datos['representante'],
            'Rafael Bermúdez Gutiérrez': datos['representante'],
            'ESCUELA HOGAR PEQUEÑOS HERMANOS': datos['institucion'],
            '08 de junio de 2026': datos['fechaExpedicion'],
            '06 de mayo de 2025': datos['fechaPeticion'],
            '1100-22-010-021': datos['claveActual'],
            '1100 22 010 021': datos['claveAnterior'],
            '1/2003': datos['predialInicio'],
            '6/2024': datos['predialFin'],
            '1/1993': datos['serviciosInicio'],
            '6/2020': datos['serviciosFin'],
            'nueve días del mes de mayo del año dos mil veintiséis': datos['fechaEjecucion'],
            '*****': datos['propietario'],
        }
        
        reemplazar_inteligente_final(doc, reemplazos)
        
        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        
        return send_file(
            doc_io,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f'Prescripcion_{datos["representante"].split()[0]}_{datetime.now().strftime("%Y%m%d")}.docx'
        )
        
    except Exception as e:
        print(f"\n❌ ERROR: {str(e)}")
        import traceback
        traceback.print_exc()
        return {'error': str(e)}, 500

@app.route('/', methods=['GET'])
def index():
    return {'status': 'OK'}, 200

if __name__ == '__main__':
    app.run(debug=True)